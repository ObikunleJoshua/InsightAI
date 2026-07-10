from pathlib import Path
from docx import Document

from services.export.exporters.base_exporter import BaseExporter


class DocxExporter(BaseExporter):
    """
    Exports reports as Microsoft Word documents.
    """

    EXPORT_FOLDER = Path("exports")

    def export(self, report: str, filename: str) -> str:

        self.EXPORT_FOLDER.mkdir(exist_ok=True)

        file_path = self.EXPORT_FOLDER / f"{filename}.docx"

        document = Document()

        document.add_heading("InsightAI Report", level=1)

        document.add_paragraph(report)

        document.save(file_path)

        return str(file_path)